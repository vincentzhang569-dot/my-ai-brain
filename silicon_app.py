import streamlit as st
from openai import OpenAI
import pdfplumber
from datetime import datetime
from io import BytesIO
import hashlib
import json
import base64
import os
from PIL import Image
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# --- 硅基流动 API 配置 ---
# Base URL 和 API Key 配置（从 Streamlit Secrets 读取）
SILICONFLOW_BASE_URL = "https://api.siliconflow.cn/v1"
SILICONFLOW_MODEL = "Qwen/Qwen2-VL-72B-Instruct"

# 从 Streamlit Secrets 读取 API Key
try:
    api_key = st.secrets["SILICONFLOW_API_KEY"]
except (FileNotFoundError, KeyError, AttributeError):
    st.error("⚠️ 未找到 API Key，请检查配置！")
    st.info("💡 请在 `.streamlit/secrets.toml` 文件中配置：\n```toml\nSILICONFLOW_API_KEY = \"sk-your-key-here\"\n```")
    api_key = None
    st.stop()

# 初始化 OpenAI 客户端（用于调用硅基流动 API）
@st.cache_resource
def get_openai_client():
    return OpenAI(
        api_key=api_key,
        base_url=SILICONFLOW_BASE_URL
    )

# --- 1. 页面配置 (移动端优先) ---
st.set_page_config(
    page_title="工业智脑 Mobile",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# --- 2. 移动端优化 CSS (原生App级体验) ---
st.markdown("""
<style>
    /* ========== 基础清理 ========== */
    #MainMenu {visibility: hidden;}
    header {visibility: hidden;}
    footer {visibility: hidden;}
    .stDeployButton {display: none;}
    
    /* ========== 移动端响应式布局 ========== */
    @media (max-width: 768px) {
        /* 主容器优化 */
        .block-container {
            padding-top: 1rem;
            padding-bottom: 6rem;
            padding-left: 1rem;
            padding-right: 1rem;
            max-width: 100%;
        }
        
        /* 标题优化 - 移动端更小 */
        .mobile-header {
            font-size: 1.5rem !important;
            margin-bottom: 0.8rem;
        }
        
        /* 聊天消息气泡优化 */
        .stChatMessage {
            border-radius: 18px;
            padding: 12px 16px;
            margin-bottom: 12px;
            box-shadow: 0 2px 8px rgba(0,0,0,0.08);
        }
        
        /* 用户消息样式 */
        .stChatMessage[data-testid="user"] {
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            margin-left: 20%;
        }
        
        /* AI消息样式 */
        .stChatMessage[data-testid="assistant"] {
            background: #f8f9fa;
            color: #333;
            margin-right: 20%;
        }
        
        /* 输入框固定在底部 */
        .stChatInput {
            position: fixed !important;
            bottom: 0 !important;
            left: 0 !important;
            right: 0 !important;
            z-index: 1000;
            background: white;
            padding: 12px 16px;
            box-shadow: 0 -2px 10px rgba(0,0,0,0.1);
            border-top: 1px solid #e0e0e0;
        }
        
        .stChatInput > div > div {
            border-radius: 25px;
            padding: 10px 20px;
            font-size: 16px; /* 防止iOS自动缩放 */
        }
        
        /* 按钮优化 - 触摸友好 */
        .stButton > button {
            width: 100%;
            height: 44px; /* iOS推荐的最小触摸目标 */
            border-radius: 12px;
            font-size: 16px;
            font-weight: 600;
            transition: all 0.3s ease;
        }
        
        .stButton > button:hover {
            transform: translateY(-2px);
            box-shadow: 0 4px 12px rgba(0,0,0,0.15);
        }
        
        /* 滑块优化 */
        .stSlider {
            padding: 10px 0;
        }
        
        /* 文件上传器优化 - 仅移动端 */
        .stFileUploader {
            padding: 15px;
            border: 2px dashed #667eea;
            border-radius: 12px;
            background: #f8f9ff;
        }
        
        /* 文件上传器标签文字优化 */
        .stFileUploader label {
            line-height: 1.5;
            word-break: break-word;
        }
        
        /* 折叠面板优化 */
        .streamlit-expanderHeader {
            font-size: 16px;
            font-weight: 600;
            padding: 12px;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            border-radius: 10px;
        }
        
        /* 文本输入框优化 */
        .stTextInput > div > div > input {
            font-size: 16px; /* 防止iOS自动缩放 */
            padding: 12px;
            border-radius: 10px;
        }
        
        /* 成功/错误消息优化 */
        .stSuccess, .stInfo, .stWarning, .stError {
            border-radius: 10px;
            padding: 12px;
            margin: 10px 0;
        }
        
        /* 列布局优化 - 移动端单列 */
        .stColumn {
            padding: 0 5px;
        }
    }
    
    /* ========== PC端保持宽屏体验 ========== */
    @media (min-width: 769px) {
        .block-container {
            padding-top: 2rem;
            padding-bottom: 2rem;
            max-width: 1200px;
        }
        
        .mobile-header {
            font-size: 2.2rem;
        }
        
        /* PC端文件上传器优化 - 防止文字重叠 */
        .stFileUploader {
            padding: 20px;
            line-height: 1.6;
            min-height: auto;
        }
        
        /* 文件上传器标签文字 */
        .stFileUploader label,
        .stFileUploader label p {
            line-height: 1.6 !important;
            white-space: normal !important;
            word-break: break-word !important;
            display: block !important;
            margin-bottom: 8px !important;
            height: auto !important;
            overflow: visible !important;
        }
        
        /* 文件上传器内部容器 */
        .stFileUploader > div {
            line-height: 1.6;
            height: auto;
        }
        
        /* 文件上传器提示文字 */
        .stFileUploader small,
        .stFileUploader .stMarkdown {
            line-height: 1.5;
            display: block;
            margin-top: 8px;
        }
        
        /* PC端文本输入框优化 */
        .stTextInput label {
            line-height: 1.5;
            white-space: normal;
        }
        
        /* PC端折叠面板内容优化 */
        .streamlit-expanderContent {
            padding: 1rem;
        }
    }
    
    /* ========== 通用样式 ========== */
    .mobile-header {
        font-weight: 800;
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
        text-align: center;
        margin-bottom: 1.5rem;
        letter-spacing: 1px;
    }
    
    /* PC端通用文字优化 - 防止重叠 */
    @media (min-width: 769px) {
        /* 所有标签文字 */
        label, label p {
            line-height: 1.5 !important;
            white-space: normal !important;
            height: auto !important;
        }
        
        /* Streamlit组件文字 */
        .stMarkdown, .stMarkdown p {
            line-height: 1.6;
            white-space: normal;
        }
    }
    
    /* 卡片式容器 */
    .card-container {
        background: white;
        border-radius: 16px;
        padding: 20px;
        margin: 15px 0;
        box-shadow: 0 4px 20px rgba(0,0,0,0.08);
    }
    
    /* 滚动条美化 */
    ::-webkit-scrollbar {
        width: 6px;
    }
    
    ::-webkit-scrollbar-track {
        background: #f1f1f1;
    }
    
    ::-webkit-scrollbar-thumb {
        background: #888;
        border-radius: 3px;
    }
    
    ::-webkit-scrollbar-thumb:hover {
        background: #555;
    }
    
    /* 加载动画优化 */
    .stSpinner > div {
        border-color: #667eea transparent transparent transparent;
    }
    
    /* 移动端触摸反馈 */
    @media (max-width: 768px) {
        * {
            -webkit-tap-highlight-color: rgba(102, 126, 234, 0.2);
        }
        
        /* 防止双击缩放 */
        * {
            touch-action: manipulation;
        }
    }
    
    /* 移动端安全区域适配 (iPhone X 等) */
    @supports (padding: max(0px)) {
        @media (max-width: 768px) {
            .stChatInput {
                padding-bottom: max(12px, env(safe-area-inset-bottom));
            }
        }
    }
    
    /* ========== 快捷指令按钮样式 ========== */
    .quick-actions-container {
        display: flex;
        gap: 10px;
        margin-bottom: 20px;
        flex-wrap: wrap;
        justify-content: center;
    }
    
    @media (max-width: 768px) {
        .quick-actions-container {
            gap: 8px;
            margin-bottom: 15px;
        }
        
        /* 移动端快捷指令按钮优化 */
        .stButton > button[kind="secondary"] {
            font-size: 13px;
            padding: 10px 12px;
        }
    }
    
    /* 下载按钮样式 */
    .download-btn-container {
        margin-top: 20px;
        padding: 15px;
        background: #f8f9fa;
        border-radius: 12px;
        text-align: center;
    }
</style>
""", unsafe_allow_html=True)

# --- 应用版本号（用于检测代码更新）---
APP_VERSION = "1.0.0"  # 更新代码时修改此版本号

# --- 注入 JavaScript 用于本地存储 ---
st.markdown(f"""
<script>
// 本地存储管理 - 保存和恢复状态
(function() {{
    const APP_VERSION = '{APP_VERSION}';
    
    // 检查版本更新
    const savedVersion = localStorage.getItem('industrial_ai_version');
    if (savedVersion && savedVersion !== APP_VERSION) {{
        // 版本更新，清除旧缓存（可选，也可以保留）
        console.log('检测到版本更新:', savedVersion, '->', APP_VERSION);
        // 如果需要清除旧缓存，取消下面的注释
        // localStorage.removeItem('industrial_ai_api_key');
        // localStorage.removeItem('industrial_ai_doc_name');
        // localStorage.removeItem('industrial_ai_doc_hash');
    }}
    localStorage.setItem('industrial_ai_version', APP_VERSION);
    
    // 保存 API Key 到 localStorage
    function saveApiKey(key) {{
        if (key && key.trim()) {{
            localStorage.setItem('industrial_ai_api_key', key);
        }}
    }}
    
    // 从 localStorage 恢复 API Key
    function restoreApiKey() {{
        return localStorage.getItem('industrial_ai_api_key') || '';
    }}
    
    // 保存文档信息
    function saveDocumentInfo(fileName, contentHash) {{
        if (fileName && contentHash) {{
            localStorage.setItem('industrial_ai_doc_name', fileName);
            localStorage.setItem('industrial_ai_doc_hash', contentHash);
        }}
    }}
    
    // 获取保存的文档信息
    function getDocumentInfo() {{
        return {{
            name: localStorage.getItem('industrial_ai_doc_name') || '',
            hash: localStorage.getItem('industrial_ai_doc_hash') || ''
        }};
    }}
    
    // 清除保存的状态
    function clearSavedState() {{
        localStorage.removeItem('industrial_ai_api_key');
        localStorage.removeItem('industrial_ai_doc_name');
        localStorage.removeItem('industrial_ai_doc_hash');
        localStorage.removeItem('industrial_ai_version');
    }}
    
    // 暴露函数到全局
    window.IndustrialAIStorage = {{
        saveApiKey: saveApiKey,
        restoreApiKey: restoreApiKey,
        saveDocumentInfo: saveDocumentInfo,
        getDocumentInfo: getDocumentInfo,
        clearSavedState: clearSavedState,
        version: APP_VERSION
    }};
    
    // 页面加载时恢复 API Key
    window.addEventListener('load', function() {{
        setTimeout(function() {{
            const apiKeyInputs = document.querySelectorAll('input[type="password"]');
            apiKeyInputs.forEach(function(input) {{
                // 恢复保存的 API Key
                const savedKey = window.IndustrialAIStorage.restoreApiKey();
                if (savedKey && !input.value) {{
                    input.value = savedKey;
                    // 触发 change 事件
                    input.dispatchEvent(new Event('change', {{ bubbles: true }}));
                    input.dispatchEvent(new Event('input', {{ bubbles: true }}));
                }}
                
                // 监听输入变化并保存
                input.addEventListener('input', function(e) {{
                    window.IndustrialAIStorage.saveApiKey(e.target.value);
                }});
            }});
        }}, 1000);
    }});
}})();
</script>
""", unsafe_allow_html=True)

# --- 核心函数：读取PDF ---
def read_pdf_text(uploaded_file) -> str:
    text_parts = []
    try:
        with pdfplumber.open(uploaded_file) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
        return "\n".join(text_parts).strip()
    except Exception as e:
        st.error(f"解析PDF出错了: {e}")
        return ""

# --- 计算文档内容哈希 ---
def calculate_content_hash(content: str) -> str:
    """计算文档内容的哈希值，用于检测文档是否已加载"""
    if not content:
        return ""
    return hashlib.md5(content.encode('utf-8')).hexdigest()

# --- 生成对话记录 Markdown ---
def generate_markdown_export(messages, doc_name=""):
    """生成Markdown格式的对话记录"""
    md_content = f"""# 工业智脑对话记录

**生成时间**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
**文档名称**: {doc_name if doc_name else '未上传文档'}

---

"""
    # 过滤掉 system 消息，只导出用户和助手的对话
    filtered_messages = [msg for msg in messages if msg.get("role") != "system"]
    for i, msg in enumerate(filtered_messages, 1):
        role = "用户" if msg["role"] == "user" else "AI助手"
        md_content += f"## {i}. {role}\n\n"
        md_content += f"{msg['content']}\n\n"
        md_content += "---\n\n"
    
    return md_content

# --- 生成对话记录 Word ---
def generate_word_export(messages, doc_name=""):
    """生成Word格式的对话记录"""
    if not DOCX_AVAILABLE:
        return None
    
    doc = Document()
    
    # 设置文档样式
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(11)
    
    # 标题
    title = doc.add_heading('工业智脑对话记录', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.name = '微软雅黑'
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    
    # 元信息
    meta_para = doc.add_paragraph()
    meta_para.add_run(f'生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}').font.name = '微软雅黑'
    meta_para = doc.add_paragraph()
    meta_para.add_run(f'文档名称: {doc_name if doc_name else "未上传文档"}').font.name = '微软雅黑'
    doc.add_paragraph('')
    
    # 分隔线
    sep_para = doc.add_paragraph('─' * 60)
    sep_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('')
    
    # 对话内容（过滤掉 system 消息）
    filtered_messages = [msg for msg in messages if msg.get("role") != "system"]
    for i, msg in enumerate(filtered_messages, 1):
        role = "用户" if msg["role"] == "user" else "AI助手"
        heading = doc.add_heading(f'{i}. {role}', level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        heading_run = heading.runs[0]
        heading_run.font.name = '微软雅黑'
        heading_run.font.size = Pt(14)
        heading_run.font.bold = True
        heading_run.font.color.rgb = RGBColor(102, 126, 234) if msg["role"] == "user" else RGBColor(0, 0, 0)
        
        # 添加内容
        content = msg['content']
        paragraphs = content.split('\n')
        for para in paragraphs:
            if para.strip():
                p = doc.add_paragraph()
                p.style.font.name = '微软雅黑'
                # 处理Markdown加粗
                parts = para.split('**')
                for idx, part in enumerate(parts):
                    run = p.add_run(part)
                    run.font.name = '微软雅黑'
                    if idx % 2 == 1:  # 奇数索引是加粗内容
                        run.bold = True
                if not parts:
                    run = p.add_run(para)
                    run.font.name = '微软雅黑'
        
        doc.add_paragraph('')
        sep_para = doc.add_paragraph('─' * 60)
        sep_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('')
    
    # 保存到BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- 初始化状态 ---
# System Prompt（工业机器人故障诊断专家）
SYSTEM_PROMPT = """你是一位资深的工业机器人故障诊断专家。请根据用户提供的文字描述或报错图片进行诊断。

输出格式要求：

1. **故障分析**：简述可能的原因。

2. **排查步骤**：分步骤列出检查点（如万用表测量哪里、检查哪根线缆）。

3. **安全警告**：提示操作风险（如高压电、机械臂坠落风险）。"""

if "messages" not in st.session_state:
    st.session_state.messages = [
        {
            "role": "system",
            "content": SYSTEM_PROMPT
        },
        {
            "role": "assistant",
            "content": "🤖 您好！我是您的工业机器人故障诊断专家。\n\n**我可以帮您：**\n1. 分析故障代码和错误信息\n2. 诊断设备故障原因\n3. 提供详细的排查步骤\n4. 提醒安全操作注意事项\n\n**使用方式：**\n• 直接描述故障现象\n• 上传故障图片（我会分析错误代码、线缆状态、仪表盘读数等）\n• 上传技术手册进行基于文档的问答\n\n💡 请开始描述您遇到的故障问题！"
        }
    ]
if "current_file" not in st.session_state:
    st.session_state.current_file = ""
if "pdf_content" not in st.session_state:
    st.session_state.pdf_content = ""
if "pending_quick_action" not in st.session_state:
    st.session_state.pending_quick_action = None
if "doc_hash" not in st.session_state:
    st.session_state.doc_hash = ""
if "restored_from_cache" not in st.session_state:
    st.session_state.restored_from_cache = False
if "selected_model" not in st.session_state:
    st.session_state.selected_model = "Qwen/Qwen2.5-7B-Instruct"
if "uploaded_image" not in st.session_state:
    st.session_state.uploaded_image = None

# 预设问题（Quick Prompts）- 工业现场快速提问
QUICK_PROMPTS = [
    "查伺服电机故障",
    "查通讯超时",
    "ABB 机器人错误代码大全",
    "编码器故障排查",
    "PLC 通讯异常"
]

# 预设问题（Quick Prompts）
QUICK_PROMPTS = [
    "查伺服电机故障",
    "查通讯超时",
    "ABB 机器人错误代码大全",
    "编码器故障排查",
    "PLC 通讯异常"
]

# --- 3. 界面布局 (移动端优化) ---

# 顶部标题 (渐变色酷炫标题)
st.markdown('<p class="mobile-header">🏭 INDUSTRIAL AI BRAIN</p>', unsafe_allow_html=True)

# === 设置面板 (移动端友好的折叠设计) ===
with st.expander("⚙️ 设置", expanded=False):
    # 状态指示器
    status_col1, status_col2, status_col3 = st.columns(3)
    with status_col1:
        st.caption(f"API Key: ✅")
    with status_col2:
        doc_status = "✅" if st.session_state.pdf_content else "⭕"
        st.caption(f"文档: {doc_status}")
    with status_col3:
        model_name = SILICONFLOW_MODEL.split('/')[-1]
        st.caption(f"模型: {model_name}")
    
    # 恢复保存的状态（仅在首次加载时）
    if not st.session_state.restored_from_cache:
        # 使用JavaScript读取localStorage并设置到session_state
        st.markdown("""
        <script>
        // 读取localStorage中的API Key
        const savedApiKey = localStorage.getItem('industrial_ai_api_key') || '';
        const savedDocInfo = {
            name: localStorage.getItem('industrial_ai_doc_name') || '',
            hash: localStorage.getItem('industrial_ai_doc_hash') || ''
        };
        
        // 将值传递给Streamlit（通过URL参数或组件通信）
        if (savedApiKey) {
            // 触发Streamlit事件来设置值
            window.parent.postMessage({
                type: 'streamlit:setComponentValue',
                key: 'restored_api_key',
                value: savedApiKey
            }, '*');
        }
        </script>
        """, unsafe_allow_html=True)
        st.session_state.restored_from_cache = True
    
    st.divider()
    
    # 1. API Key 状态显示
    st.success(f"✅ API Key 已配置（前4位: {api_key[:4]}...）")
    
    # 显示模型信息
    st.markdown("**🧠 当前使用的模型**")
    st.info("**Qwen2-VL-72B** (720亿参数)")
    st.caption("强大的视觉语言模型，支持图片分析和故障诊断")
    
    st.divider()
    
    # 一键重置对话按钮
    if st.button("🗑️ 开启新对话", use_container_width=True, help="清空当前对话历史，开始新的对话"):
        st.session_state.messages = [
            {
                "role": "system",
                "content": SYSTEM_PROMPT
            },
            {
                "role": "assistant",
                "content": "🤖 对话已重置，请开始新的故障诊断咨询。"
            }
        ]
        st.session_state.uploaded_image = None
        st.rerun()
    
    st.divider()
    
    # 2. 文件上传 (移动端优化)
    st.markdown("**📄 上传技术手册**")
    uploaded_file = st.file_uploader(
        "支持 PDF 格式", 
        type=["pdf"],
        help="上传您的技术手册、操作指南或故障排除文档"
    )
    
    # 显示文档信息
    if st.session_state.pdf_content:
        doc_length = len(st.session_state.pdf_content)
        st.success(f"📚 当前文档: **{st.session_state.current_file}**")
        st.caption(f"📊 文档大小: {doc_length:,} 字符 | 约 {doc_length//1000}K tokens")
    
    st.divider()
    
    # 3. 图片上传组件（侧边栏）
    st.markdown("**📷 上传故障图片**")
    uploaded_image = st.file_uploader(
        "支持 PNG、JPG 格式",
        type=['png', 'jpg', 'jpeg'],
        help="上传设备故障照片，AI 会分析图片中的错误代码、线缆状态或仪表盘读数",
        key="sidebar_image_uploader"
    )
    
    # 如果上传了图片，显示缩略图
    if uploaded_image is not None:
        try:
            image = Image.open(uploaded_image)
            st.image(image, caption="上传的图片", use_container_width=True)
            st.session_state.uploaded_image = uploaded_image
            st.success("✅ 图片已上传")
        except Exception as e:
            st.error(f"❌ 图片处理失败: {str(e)}")
            st.session_state.uploaded_image = None
    else:
        st.session_state.uploaded_image = None
    
    st.divider()
    
    # 4. 高级设置 (移动端紧凑布局)
    st.markdown("**⚙️ 高级设置**")
    col_p1, col_p2 = st.columns([2, 1])
    with col_p1:
        temperature = st.slider(
            "🎨 创造力 (Temperature)", 
            0.0, 1.0, 0.3,
            help="数值越低越严谨，数值越高越有创造性"
        )
    with col_p2:
        st.markdown("<br>", unsafe_allow_html=True)  # 对齐按钮
        if st.button("🗑️ 清空", use_container_width=True, help="清空所有对话记录"):
            st.session_state.messages = [{"role": "assistant", "content": "🤖 对话历史已清空，请重新提问。"}]
            st.rerun()
    
    # 清除缓存按钮
    st.markdown("---")
    if st.button("🗑️ 清除所有保存的状态（API Key和文档信息）", use_container_width=True, help="清除浏览器中保存的所有缓存数据"):
        st.markdown("""
        <script>
        if (window.IndustrialAIStorage) {
            window.IndustrialAIStorage.clearSavedState();
            alert('缓存已清除！');
        }
        </script>
        """, unsafe_allow_html=True)
        # 清除session_state
        # API Key 现在从 secrets 读取，无需清除
        st.session_state.pdf_content = ""
        st.session_state.current_file = ""
        st.session_state.doc_hash = ""
        st.session_state.messages = [{"role": "assistant", "content": "🤖 所有状态已清除，请重新配置。"}]
        st.rerun()

    # 处理文件读取
    if uploaded_file:
        if st.session_state.current_file != uploaded_file.name:
            with st.spinner("📄 正在解析文档，请稍候..."):
                text = read_pdf_text(uploaded_file)
                st.session_state.pdf_content = text
                st.session_state.current_file = uploaded_file.name
                # 计算文档哈希并保存到localStorage
                doc_hash = calculate_content_hash(text)
                st.session_state.doc_hash = doc_hash
                # 保存文档信息到localStorage
                st.markdown(f"""
                <script>
                if (window.IndustrialAIStorage) {{
                    window.IndustrialAIStorage.saveDocumentInfo('{uploaded_file.name}', '{doc_hash}');
                }}
                </script>
                """, unsafe_allow_html=True)
            st.success(f"✅ 文档加载成功: **{uploaded_file.name}**")
            st.info("💾 状态已自动保存，刷新页面后不会丢失（API Key和文档信息）")
            st.balloons()  # 成功提示动画

# --- 4. 聊天区域 (移动端优化) ---

# 动态状态提示
if st.session_state.pdf_content:
    st.success("✅ 一切就绪！您可以基于文档提问，也可以直接提问任何问题。")
    st.caption("💾 您的配置已自动保存，刷新页面后不会丢失")
else:
    st.info("💡 已就绪！您可以开始提问。如需基于文档问答，可在设置中上传PDF文档。")

# 添加使用提示
    st.markdown("""
    <div style="background: #e3f2fd; padding: 12px; border-radius: 8px; margin: 10px 0;">
        <strong>📌 使用提示：</strong><br>
        • 本网站不局限于工业问题
        • 您可以先输入API Key直接提问，也可以上传文档后进行基于文档的问答<br>
        • 支持技术咨询、代码问题、学习辅导、生活问答等多种场景
    </div>
    """, unsafe_allow_html=True)

# === 预设问题按钮 (Quick Prompts) - 工业现场快速提问 ===
st.markdown("**⚡ 快速提问（点击下方按钮）**")

# 使用列布局显示预设问题按钮
# PC端：5个按钮并排；移动端：自动换行
prompt_cols = st.columns(5)
for idx, prompt_text in enumerate(QUICK_PROMPTS):
    with prompt_cols[idx]:
        if st.button(
            prompt_text, 
            key=f"quick_prompt_{idx}",
            use_container_width=True,
            help=f"快速提问：{prompt_text}"
        ):
            st.session_state.pending_quick_action = prompt_text
            st.rerun()

st.markdown("---")

# === 快捷指令按钮 (工业现场一键操作) ===
# 只有在有文档时才显示快捷指令
if st.session_state.pdf_content:
    st.markdown("**⚡ 快捷指令**")
    col1, col2, col3 = st.columns(3)
    
    with col1:
        if st.button("📄 全文摘要", use_container_width=True, help="一键生成文档全文摘要"):
            st.session_state.pending_quick_action = "请为这份文档生成一份详细的全文摘要，包括主要章节、核心内容和关键要点。"
            st.rerun()
    
    with col2:
        if st.button("🔧 故障诊断", use_container_width=True, help="快速进入故障排查模式"):
            st.session_state.pending_quick_action = "请列出这份文档中涉及的所有故障代码、故障原因和对应的解决方案。如果文档中没有相关内容，请说明。"
            st.rerun()
    
    with col3:
        if st.button("⚠️ 安全须知", use_container_width=True, help="查看安全操作注意事项"):
            st.session_state.pending_quick_action = "请提取这份文档中所有关于安全操作、注意事项、警告信息的内容，并按重要性排序。"
            st.rerun()

# 显示聊天记录 (移动端优化)
for msg in st.session_state.messages:
    # 跳过 system 消息，不显示给用户
    if msg.get("role") == "system":
        continue
    # 只显示 user 和 assistant 的消息
    with st.chat_message(msg["role"]):
        st.markdown(msg["content"])

# === 对话记录下载功能 ===
if len(st.session_state.messages) > 1:  # 至少有用户和AI的对话
    st.divider()
    st.markdown("**💾 导出对话记录**")
    
    download_col1, download_col2 = st.columns(2)
    
    with download_col1:
        # Markdown下载
        md_content = generate_markdown_export(
            st.session_state.messages, 
            st.session_state.current_file
        )
        st.download_button(
            label="📄 下载 Markdown",
            data=md_content,
            file_name=f"对话记录_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md",
            mime="text/markdown",
            use_container_width=True,
            help="下载为Markdown格式，适合文档管理"
        )
    
    with download_col2:
        # Word下载
        if DOCX_AVAILABLE:
            word_buffer = generate_word_export(
                st.session_state.messages,
                st.session_state.current_file
            )
            if word_buffer:
                st.download_button(
                    label="📝 下载 Word",
                    data=word_buffer,
                    file_name=f"对话记录_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    help="下载为Word格式，适合工单汇报"
                )
        else:
            st.info("💡 安装 python-docx 以支持Word导出\n`pip install python-docx`", icon="ℹ️")

# --- 5. 图片处理函数 ---
def image_to_base64(image):
    """将 PIL Image 对象转换为 Base64 字符串"""
    buffered = BytesIO()
    # 转换为 RGB 模式（如果是 RGBA 等）
    if image.mode != 'RGB':
        image = image.convert('RGB')
    # 保存为 PNG 格式到内存
    image.save(buffered, format="PNG")
    # 转换为 Base64
    img_base64 = base64.b64encode(buffered.getvalue()).decode('utf-8')
    return img_base64

# --- 6. 处理用户输入 (移动端优化) ---
# 处理快捷指令
prompt = None
if st.session_state.pending_quick_action:
    prompt = st.session_state.pending_quick_action
    st.session_state.pending_quick_action = None  # 清除标志

# 处理用户手动输入
user_input = st.chat_input("💬 输入技术问题...")
if user_input:
    prompt = user_input

if prompt:
    
    # 检查是否有图片
    has_image = st.session_state.uploaded_image is not None
    
    # 将图片转换为 Base64（如果上传了图片）
    image_base64 = None
    if has_image:
        try:
            image = Image.open(st.session_state.uploaded_image)
            image_base64 = image_to_base64(image)
        except Exception as e:
            st.error(f"❌ 图片处理失败: {str(e)}")
            has_image = False
            image_base64 = None
    
    # 添加用户消息
    user_message_content = prompt
    if has_image:
        user_message_content = f"[包含图片] {prompt}"
    st.session_state.messages.append({"role": "user", "content": user_message_content})
    with st.chat_message("user"):
        st.markdown(prompt)
        if has_image and st.session_state.uploaded_image:
            st.image(st.session_state.uploaded_image, use_container_width=True)

    # 构建系统提示词（根据是否有文档选择不同策略）
    pdf_text = st.session_state.pdf_content
    
    # 使用基础 System Prompt，并根据文档情况增强
    if pdf_text:
        # 有文档：基于文档回答
        system_prompt = f"""{SYSTEM_PROMPT}

【文档分析模式】：
- 严格基于以下文档内容回答问题，不要编造信息
- 如果问题超出文档范围，明确告知"文档中未提及此内容"，但可以基于你的知识提供一般性建议
- 结合文档内容和图片分析（如有）给出综合诊断

【文档内容】：
{pdf_text[:8000]}

请严格按照输出格式要求回答用户问题。"""
    else:
        # 无文档：使用基础 System Prompt
        system_prompt = SYSTEM_PROMPT

    # 调用硅基流动 API
    try:
        client = get_openai_client()
        
        with st.chat_message("assistant"):
            # 显示加载状态
            model_display_name = SILICONFLOW_MODEL.split('/')[-1]
            spinner_text = f"🤔 {model_display_name} 正在分析中..."
            if has_image:
                spinner_text = f"👁️ {model_display_name} 正在分析图片..."
            
            with st.spinner(spinner_text):
                # 构建用户消息内容（OpenAI 标准格式）
                user_content = []
                
                # 添加文本
                user_content.append({"type": "text", "text": prompt})
                
                # 如果有图片，添加图片（Base64 格式）
                if has_image and image_base64:
                    user_content.append({
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:image/png;base64,{image_base64}"
                        }
                    })
                
                # 构建消息列表（OpenAI 标准格式）
                messages = [{"role": "system", "content": system_prompt}]
                
                # 添加历史消息（过滤掉 system 消息，因为我们已经添加了新的 system prompt）
                for msg in st.session_state.messages[:-1]:
                    # 跳过 system 消息
                    if msg.get("role") == "system":
                        continue
                    # 历史消息可能是简单格式或复杂格式
                    if isinstance(msg.get("content"), list):
                        # 如果已经是复杂格式，直接添加
                        messages.append(msg)
                    else:
                        # 如果是简单格式，转换为标准格式
                        messages.append({
                            "role": msg["role"],
                            "content": msg["content"]
                        })
                
                # 添加当前用户消息
                messages.append({
                    "role": "user",
                    "content": user_content
                })
                
                # 调用硅基流动 API（流式输出）
                stream = client.chat.completions.create(
                    model=SILICONFLOW_MODEL,
                    messages=messages,
                    stream=True
                )
                
                # 处理流式响应
                response = st.write_stream(stream)
        
        # 保存回复
        st.session_state.messages.append({"role": "assistant", "content": response})
        
        # 反馈机制（隐形埋点）
        feedback_key = f"feedback_{len(st.session_state.messages)}"
        
        # 检查是否已经反馈过
        if feedback_key not in st.session_state.get('feedback_given', {}):
            # 使用按钮实现反馈功能（兼容性更好）
            st.markdown("**这个回答对您有帮助吗？**")
            feedback_col1, feedback_col2 = st.columns(2)
            feedback = None
            
            with feedback_col1:
                if st.button("👍 有帮助", key=f"{feedback_key}_positive", use_container_width=True):
                    feedback = "positive"
            
            with feedback_col2:
                if st.button("👎 没帮助", key=f"{feedback_key}_negative", use_container_width=True):
                    feedback = "negative"
            
            # 记录反馈（静默记录，不打扰用户）
            if feedback:
                # 标记已反馈，避免重复记录
                if 'feedback_given' not in st.session_state:
                    st.session_state.feedback_given = {}
                st.session_state.feedback_given[feedback_key] = True
                
                # 初始化反馈数据存储
                if 'feedback_data' not in st.session_state:
                    st.session_state.feedback_data = []
                
                # 记录反馈数据
                feedback_entry = {
                    'query': prompt,
                    'response': response[:200] if len(response) > 200 else response,  # 截断长响应
                    'feedback': feedback,
                    'has_image': has_image,
                    'has_document': bool(pdf_text),
                    'model': SILICONFLOW_MODEL,
                    'timestamp': len(st.session_state.messages)
                }
                st.session_state.feedback_data.append(feedback_entry)
                
                # 显示感谢提示
                st.success("✅ 感谢您的反馈！")
                
                # 可以在这里添加：
                # 1. 保存到本地文件（JSON/CSV）
                # 2. 发送到分析平台
                # 3. 保存到数据库
                # 示例：保存到本地 JSON 文件（可选，取消注释以启用）
                # try:
                #     import json
                #     with open('feedback_log.json', 'a', encoding='utf-8') as f:
                #         json.dump(feedback_entry, f, ensure_ascii=False)
                #         f.write('\n')
                # except:
                #     pass
        else:
            # 已经反馈过，显示已反馈状态
            st.caption("💡 您已对此回答进行反馈，感谢！")
        
        # 清除图片状态（可选：如果希望图片保留在对话中，可以注释掉）
        # st.session_state.image_base64 = None
        # st.session_state.uploaded_image = None
        
        # 成功提示
        st.toast("✅ 回答完成", icon="✅")

    except Exception as e:
        error_msg = str(e)
        st.error(f"❌ 请求失败: {error_msg}")
        st.info("💡 请检查：1) API Key 是否正确 2) 网络连接是否正常 3) 账户余额是否充足 4) 模型是否支持图片输入")
        
        # 添加重试建议
        if "401" in error_msg or "Unauthorized" in error_msg:
            st.warning("🔑 API Key 验证失败，请检查密钥是否正确")
        elif "429" in error_msg or "rate limit" in error_msg.lower():
            st.warning("⏱️ 请求过于频繁，请稍后再试")
        elif has_image and ("vision" in error_msg.lower() or "image" in error_msg.lower()):
            st.warning("📷 当前模型可能不支持图片输入。如需图片分析功能，请使用支持视觉的模型（如 Qwen2.5-VL-7B-Instruct）。")

